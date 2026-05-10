"""Agent 2: 多格式文档文本提取。"""
import time

from utils import (
    anonymize,
    iter_ipynb_cells,
    iter_pptx_content,
    read_text_file,
    save_manifest,
    setup_logging,
)

TEXT_MAX_CHARS = 50000


class ExtractorAgent:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        self.logger = setup_logging(output_dir)
        self._extractors = {
            ".docx": self._extract_docx,
            ".pptx": self._extract_pptx,
            ".pdf": self._extract_pdf,
            ".xlsx": self._extract_xlsx,
            ".ipynb": self._extract_ipynb,
            ".csv": self._extract_csv,
            ".txt": self._extract_text,
            ".md": self._extract_text,
        }

    def run(self, manifest):
        self.logger.info("=" * 50)
        self.logger.info("[EXTRACTOR] 开始文本提取")
        t0 = time.time()
        ok = fail = total_chars = 0

        for f in manifest.get("files", []):
            result = {"text_content": "", "text_length": 0, "extraction_ok": False, "error_message": None}
            if f.get("is_duplicate"):
                result["text_content"] = "[重复文件，跳过提取]"
                result["error_message"] = "duplicate"
            else:
                try:
                    text = self._extract(f)
                    if text:
                        text = anonymize(text)
                        result["text_length"] = len(text)
                        result["text_content"] = text[:TEXT_MAX_CHARS]
                        result["extraction_ok"] = True
                        total_chars += len(text)
                        ok += 1
                    else:
                        result["extraction_ok"] = True
                        ok += 1
                except Exception as e:
                    result["error_message"] = str(e)
                    fail += 1
                    self.logger.warning("[EXTRACTOR] FAIL: %s — %s", f["filename"], e)
            f.update({k: v for k, v in result.items() if k in f or k in result})

        manifest["stats"]["extraction_ok"] = ok
        manifest["stats"]["extraction_fail"] = fail
        manifest["stats"]["total_chars"] = total_chars
        save_manifest(manifest, self.output_dir)
        manifest["stats"]["extractor_time"] = round(time.time() - t0, 2)
        self.logger.info("[EXTRACTOR] 完成 — %d OK, %d FAIL, %d 字符", ok, fail, total_chars)
        return manifest

    def _extract(self, f):
        ext = f["ext"]
        handler = self._extractors.get(ext)
        if handler:
            return handler(f["path"])
        return ""

    def _extract_docx(self, path):
        import docx
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)

    def _extract_pptx(self, path):
        parts = []
        for i, paragraphs, table_rows in iter_pptx_content(path):
            lines = [f"--- Slide {i} ---"] + paragraphs
            if table_rows:
                lines.append(" | ".join(table_rows[0]))
            parts.append("\n".join(lines))
        return "\n\n".join(parts)

    def _extract_pdf(self, path):
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _extract_xlsx(self, path):
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for name in wb.sheetnames:
            ws = wb[name]
            lines = [f"--- Sheet: {name} ---"]
            for row in ws.iter_rows(max_row=500):
                line = "\t".join(str(cell.value) if cell.value is not None else "" for cell in row).strip()
                if line:
                    lines.append(line)
            parts.append("\n".join(lines))
        wb.close()
        return "\n\n".join(parts)

    def _extract_ipynb(self, path):
        parts = []
        for i, cell_type, source in iter_ipynb_cells(path):
            if cell_type == "code":
                parts.append(f"[In {i + 1}]:\n{source}")
            elif cell_type == "markdown":
                parts.append(f"[Markdown {i + 1}]:\n{source}")
        return "\n\n".join(parts)

    def _extract_csv(self, path):
        return read_text_file(path)

    def _extract_text(self, path):
        return read_text_file(path)
